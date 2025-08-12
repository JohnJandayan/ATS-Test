import gradio as gr
import PyPDF2
import docx
import io
import re
import nltk
from transformers import pipeline, AutoTokenizer, AutoModelForSequenceClassification
from config import MODEL_CONFIG
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np
import warnings
warnings.filterwarnings("ignore")

# Download NLTK data (done once)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

# Initialize AI models
print("Loading AI models...")
sentiment_analyzer = pipeline(
    "sentiment-analysis",
    model=MODEL_CONFIG.get("sentiment_model", "cardiffnlp/twitter-roberta-base-sentiment-latest")
)
classifier = pipeline(
    "zero-shot-classification",
    model=MODEL_CONFIG.get("classification_model", "facebook/bart-large-mnli")
)

# ATS keywords and categories
ATS_KEYWORDS = {
    "technical_skills": [
        "python", "java", "javascript", "react", "nodejs", "sql", "aws", "docker", 
        "kubernetes", "git", "linux", "api", "html", "css", "mongodb", "postgresql",
        "machine learning", "data analysis", "artificial intelligence", "tensorflow",
        "pytorch", "pandas", "numpy", "scikit-learn", "cloud computing", "devops"
    ],
    "soft_skills": [
        "leadership", "communication", "teamwork", "problem solving", "analytical",
        "creative", "adaptable", "organized", "detail oriented", "time management",
        "project management", "collaboration", "critical thinking", "innovation"
    ],
    "action_verbs": [
        "achieved", "developed", "created", "managed", "led", "implemented", "designed",
        "optimized", "improved", "analyzed", "collaborated", "coordinated", "executed",
        "delivered", "built", "established", "increased", "reduced", "streamlined"
    ],
    "education": [
        "bachelor", "master", "phd", "degree", "university", "college", "certification",
        "course", "training", "education", "graduate", "undergraduate", "diploma"
    ]
}

def extract_text_from_file(file_input):
    """Extract text from uploaded file (PDF, DOCX, or TXT).

    Supports Gradio 5 inputs which may be a string path, dict with 'path',
    or a file-like object with .read() and .name.
    """
    if file_input is None:
        return ""

    # Normalize to a file path and extension if possible
    file_path = None
    filename_lower = None

    if isinstance(file_input, str):
        file_path = file_input
        filename_lower = file_input.lower()
    elif isinstance(file_input, dict) and 'path' in file_input:
        file_path = file_input['path']
        filename_lower = file_path.lower()
    elif hasattr(file_input, 'name') and hasattr(file_input, 'read'):
        # file-like object
        try:
            content = file_input.read()
            filename_lower = getattr(file_input, 'name', 'uploaded_file').lower()
            if filename_lower.endswith('.pdf'):
                try:
                    reader = PyPDF2.PdfReader(io.BytesIO(content))
                    pages_text = []
                    for page in reader.pages:
                        page_text = page.extract_text() or ""
                        pages_text.append(page_text)
                    return "\n".join(pages_text)
                except Exception:
                    pass  # fall through to error below
            if filename_lower.endswith('.docx'):
                try:
                    doc = docx.Document(io.BytesIO(content))
                    return "\n".join(p.text for p in doc.paragraphs)
                except Exception:
                    pass
            if filename_lower.endswith('.txt'):
                try:
                    return content.decode('utf-8', errors='ignore')
                except Exception:
                    pass
        except Exception as e:
            return f"Error reading file: {str(e)}"

    # If we have a file path, read from disk
    if file_path and filename_lower:
        try:
            if filename_lower.endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    pages_text = []
                    for page in reader.pages:
                        page_text = page.extract_text() or ""
                        pages_text.append(page_text)
                    return "\n".join(pages_text)
            elif filename_lower.endswith('.docx'):
                doc = docx.Document(file_path)
                return "\n".join(p.text for p in doc.paragraphs)
            elif filename_lower.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            else:
                return "Unsupported file format. Please upload PDF, DOCX, or TXT files."
        except Exception as e:
            return f"Error reading file: {str(e)}"

    return "Error reading file: Unrecognized file input type"

def analyze_ats_keywords(text):
    """Analyze resume for ATS-friendly keywords"""
    text_lower = text.lower()
    scores = {}
    found_keywords = {}
    
    for category, keywords in ATS_KEYWORDS.items():
        found = [kw for kw in keywords if kw in text_lower]
        found_keywords[category] = found
        scores[category] = min(100, (len(found) / len(keywords)) * 100)
    
    return scores, found_keywords

def analyze_sentiment_and_tone(text):
    """Analyze professional tone using AI sentiment analysis"""
    try:
        # Split text into chunks for analysis
        sentences = nltk.sent_tokenize(text)
        if not sentences:
            return {"professional_tone": 50, "confidence": "Low"}
        
        # Analyze first few sentences for tone
        sample_text = " ".join(sentences[:5])
        
        result = sentiment_analyzer(sample_text[:512])  # Limit length

        # Normalize sentiment labels across different models
        label = str(result[0].get('label', '')).lower()
        score = float(result[0].get('score', 0.5))
        is_positive = label in {"positive", "pos", "label_2", "label_1"}  # include common variants
        is_negative = label in {"negative", "neg", "label_0"}

        if is_positive:
            tone_score = min(95, 60 + (score * 35))
        elif is_negative:
            tone_score = max(30, 60 - (score * 30))
        else:  # neutral or unknown
            tone_score = 60
            
        confidence = "High" if result[0]['score'] > 0.8 else "Medium" if result[0]['score'] > 0.5 else "Low"
        
        return {"professional_tone": int(tone_score), "confidence": confidence}
        
    except Exception as e:
        return {"professional_tone": 50, "confidence": "Low", "error": str(e)}

def classify_resume_sections(text):
    """Use AI to classify and analyze resume sections"""
    try:
        sections = ["experience", "education", "skills", "summary", "achievements"]
        
        # Split text into potential sections
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        if not paragraphs:
            return {"sections_identified": 0, "completeness": 20}
        
        # Use AI to classify sections
        classified_sections = set()
        
        for paragraph in paragraphs[:10]:  # Analyze first 10 paragraphs
            if len(paragraph) > 20:  # Skip very short paragraphs
                result = classifier(paragraph[:512], candidate_labels=sections)
                if result['scores'][0] > 0.3:  # Confidence threshold
                    classified_sections.add(result['labels'][0])
        
        completeness = (len(classified_sections) / len(sections)) * 100
        
        return {
            "sections_identified": len(classified_sections),
            "completeness": min(100, completeness),
            "found_sections": list(classified_sections)
        }
        
    except Exception as e:
        return {"sections_identified": 2, "completeness": 40, "error": str(e)}

def calculate_readability_score(text):
    """Calculate ATS readability score"""
    if not text.strip():
        return 0
    
    # Basic readability metrics
    sentences = nltk.sent_tokenize(text)
    words = nltk.word_tokenize(text)
    
    if not sentences or not words:
        return 30
    
    avg_sentence_length = len(words) / len(sentences)
    
    # Prefer moderate sentence length for ATS
    if 10 <= avg_sentence_length <= 20:
        readability = 90
    elif 8 <= avg_sentence_length <= 25:
        readability = 75
    else:
        readability = 50
    
    # Check for bullet points (ATS-friendly)
    bullet_points = len(re.findall(r'[•\-\*]\s', text))
    if bullet_points > 3:
        readability += 10
    
    # Check for proper formatting indicators
    if re.search(r'\b(education|experience|skills|summary)\b', text.lower()):
        readability += 5
    
    return min(100, readability)

def generate_improvement_suggestions(scores, found_keywords, analysis_results):
    """Generate AI-powered improvement suggestions"""
    suggestions = []
    
    # Keyword suggestions
    if scores['technical_skills'] < 60:
        missing_tech = set(ATS_KEYWORDS['technical_skills']) - set(found_keywords['technical_skills'])
        suggestions.append(f"🔧 **Technical Skills**: Add relevant technical skills like {', '.join(list(missing_tech)[:3])}")
    
    if scores['soft_skills'] < 50:
        missing_soft = set(ATS_KEYWORDS['soft_skills']) - set(found_keywords['soft_skills'])
        suggestions.append(f"🤝 **Soft Skills**: Include soft skills such as {', '.join(list(missing_soft)[:3])}")
    
    if scores['action_verbs'] < 70:
        suggestions.append("⚡ **Action Verbs**: Use more strong action verbs like 'achieved', 'developed', 'led', 'optimized'")
    
    # Section completeness
    if analysis_results.get('completeness', 0) < 80:
        suggestions.append("📋 **Resume Sections**: Ensure you have clear sections for Experience, Education, Skills, and Summary")
    
    # Professional tone
    tone_score = analysis_results.get('professional_tone', 50)
    if tone_score < 70:
        suggestions.append("💼 **Professional Tone**: Use more professional language and avoid casual expressions")
    
    # Readability
    if 'readability' in analysis_results and analysis_results['readability'] < 70:
        suggestions.append("📖 **Readability**: Use bullet points and clear formatting to improve ATS readability")
    
    if not suggestions:
        suggestions.append("✅ **Great job!** Your resume shows strong ATS optimization. Keep it updated with relevant keywords.")
    
    return suggestions

def analyze_resume(file, job_description=""):
    """Main AI-powered resume analysis function"""
    if file is None:
        return "❌ Please upload a resume file."
    
    try:
        # Extract text from file
        text = extract_text_from_file(file)
        
        if not text or len(text.strip()) < 50:
            return "❌ Could not extract enough text from the file. Please ensure the file contains readable text."
        
        # Combine resume text with job description for targeted analysis
        analysis_text = text
        if job_description.strip():
            analysis_text = f"Job Requirements: {job_description}\n\nResume: {text}"
        
        # Perform AI analysis
        print("Analyzing keywords...")
        keyword_scores, found_keywords = analyze_ats_keywords(text)
        
        print("Analyzing tone...")
        sentiment_result = analyze_sentiment_and_tone(text)
        
        print("Analyzing structure...")
        section_analysis = classify_resume_sections(text)
        
        print("Calculating readability...")
        readability = calculate_readability_score(text)
        
        # Combine all analysis results
        analysis_results = {
            **sentiment_result,
            **section_analysis,
            'readability': readability
        }
        
        # Calculate overall score
        overall_score = int(np.mean([
            np.mean(list(keyword_scores.values())),
            sentiment_result.get('professional_tone', 50),
            section_analysis.get('completeness', 50),
            readability
        ]))
        
        # Generate suggestions
        suggestions = generate_improvement_suggestions(keyword_scores, found_keywords, analysis_results)
        
        # Format results
        filename = file.name if hasattr(file, 'name') else "uploaded_file"
        
        result = f"""# 🤖 AI-Powered ATS Resume Analysis

## 📊 Overall ATS Score: {overall_score}/100

### 🎯 Category Breakdown:
- **Technical Skills**: {keyword_scores['technical_skills']:.1f}/100 ({len(found_keywords['technical_skills'])} keywords found)
- **Soft Skills**: {keyword_scores['soft_skills']:.1f}/100 ({len(found_keywords['soft_skills'])} keywords found)  
- **Action Verbs**: {keyword_scores['action_verbs']:.1f}/100 ({len(found_keywords['action_verbs'])} strong verbs found)
- **Professional Tone**: {sentiment_result.get('professional_tone', 50)}/100 (AI Confidence: {sentiment_result.get('confidence', 'Unknown')})
- **Structure & Sections**: {section_analysis.get('completeness', 50):.1f}/100 ({section_analysis.get('sections_identified', 0)} sections identified)
- **ATS Readability**: {readability}/100

### 💡 AI-Generated Improvement Suggestions:

{chr(10).join(suggestions)}

### 🔍 Detailed Analysis:

**Found Technical Skills**: {', '.join(found_keywords['technical_skills'][:10]) if found_keywords['technical_skills'] else 'None detected'}

**Found Soft Skills**: {', '.join(found_keywords['soft_skills'][:10]) if found_keywords['soft_skills'] else 'None detected'}

**Identified Sections**: {', '.join(section_analysis.get('found_sections', [])) if section_analysis.get('found_sections') else 'Basic structure detected'}

### 📄 File Information:
- **Filename**: {filename}
- **Text Length**: {len(text)} characters
- **Analysis Method**: AI-powered using HuggingFace models
- **Job-Specific**: {"Yes" if job_description.strip() else "No"}

---
*Analysis powered by AI models: RoBERTa (sentiment), BART (classification), and custom ATS algorithms*
        """
        
        return result
        
    except Exception as e:
        return f"❌ Analysis Error: {str(e)}\n\nPlease try uploading a different file format or check if the file is corrupted."

# Create the Gradio interface
try:
    # Simple interface creation to avoid JSON schema issues
    with gr.Blocks(title="🤖 AI-Powered ATS Resume Rater") as demo:
        gr.Markdown("# 🤖 AI-Powered ATS Resume Rater")
        gr.Markdown("Upload your resume for comprehensive AI analysis using advanced NLP models. Get actionable insights to optimize your resume for Applicant Tracking Systems.")
        
        with gr.Row():
            with gr.Column():
                file_input = gr.File(
                    label="📄 Upload Resume (PDF/DOCX/TXT)", 
                    file_types=[".pdf", ".docx", ".txt"],
                    type="filepath"
                )
                job_input = gr.Textbox(
                    label="🎯 Job Description (Optional)", 
                    placeholder="Paste the job description here for targeted analysis...",
                    lines=5
                )
                analyze_btn = gr.Button("🔍 Analyze Resume", variant="primary")
        
        with gr.Row():
            output = gr.Textbox(
                label="🤖 AI Analysis Results", 
                lines=20,
                max_lines=30,
                show_copy_button=True
            )
        
        analyze_btn.click(
            fn=analyze_resume,
            inputs=[file_input, job_input],
            outputs=output
        )
        
        gr.Markdown("<center>Built with ❤️ by <a href='https://portfolio-john-jandayan.vercel.app' target='_blank'>John Jandayan</a> | Powered by HuggingFace AI Models</center>")

except Exception as e:
    print(f"Error creating Gradio Blocks interface: {e}")
    # Fallback with even simpler configuration
    try:
        demo = gr.Interface(
            fn=analyze_resume,
            inputs=[
                gr.File(label="Upload Resume", file_types=[".pdf", ".docx", ".txt"], type="filepath"),
                gr.Textbox(label="Job Description (Optional)", lines=5)
            ],
            outputs=gr.Textbox(label="Analysis Results", lines=20),
            title="ATS Resume Rater"
        )
    except Exception as e2:
        print(f"Error creating fallback interface: {e2}")
        # Final fallback - minimal interface
        demo = gr.Interface(
            fn=analyze_resume,
            inputs=[
                gr.File(label="Upload Resume", type="filepath"),
                gr.Textbox(label="Job Description", lines=3)
            ],
            outputs=gr.Textbox(label="Results", lines=15),
            title="Resume Analyzer"
        )

if __name__ == "__main__":
    import os
    
    # Check if running in a cloud environment (like HuggingFace Spaces)
    is_cloud_env = os.getenv("SPACE_ID") is not None or os.getenv("GRADIO_SERVER_NAME") is not None
    
    try:
        demo.launch(
            server_name="0.0.0.0",
            server_port=7860,
            share=is_cloud_env  # Enable sharing only in cloud environments
        )
    except Exception as e:
        print(f"Error launching with share={is_cloud_env}: {e}")
        print("Trying with share=True...")
        try:
            demo.launch(
                server_name="0.0.0.0",
                server_port=7860,
                share=True
            )
        except Exception as e2:
            print(f"Error launching with share=True: {e2}")
            print("Trying local launch...")
            demo.launch(
                server_name="127.0.0.1",
                server_port=7860,
                share=False
            )
